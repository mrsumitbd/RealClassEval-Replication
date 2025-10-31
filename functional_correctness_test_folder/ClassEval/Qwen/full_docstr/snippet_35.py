
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path
        self.doc = Document(
            self.file_path) if self._file_exists() else Document()

    def _file_exists(self):
        try:
            with open(self.file_path, 'r'):
                return True
        except FileNotFoundError:
            return False

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        full_text = []
        for para in self.doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            paragraph = self.doc.add_paragraph(content)
            paragraph.style.font.size = Pt(font_size)
            paragraph.alignment = self._get_alignment_value(alignment)
            self.doc.save(self.file_path)
            return True
        except Exception:
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            self.doc.add_heading(heading, level=level)
            self.doc.save(self.file_path)
            return True
        except Exception:
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            rows = len(data)
            cols = len(data[0]) if rows > 0 else 0
            table = self.doc.add_table(rows=rows, cols=cols)
            for i, row in enumerate(data):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = str(cell)
            self.doc.save(self.file_path)
            return True
        except Exception:
            return False

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }
        return alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
