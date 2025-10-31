
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        try:
            if not os.path.exists(self.file_path):
                return ""
            doc = Document(self.file_path)
            parts = []
            for p in doc.paragraphs:
                parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        except Exception:
            return ""

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run(str(content) if content is not None else "")
            try:
                run.font.size = Pt(int(font_size))
            except Exception:
                run.font.size = Pt(12)
            para.alignment = self._get_alignment_value(alignment)
            doc.save(self.file_path)
            return True
        except Exception:
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            doc = self._load_or_create_document()
            lvl = int(level) if isinstance(level, (int, str)) else 1
            if lvl < 0:
                lvl = 1
            doc.add_heading(heading if heading is not None else "", level=lvl)
            doc.save(self.file_path)
            return True
        except Exception:
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            if not isinstance(data, list) or (len(data) > 0 and not all(isinstance(row, list) for row in data)):
                return False
            rows = len(data)
            cols = max((len(row) for row in data), default=0)
            if rows == 0 or cols == 0:
                return False

            doc = self._load_or_create_document()
            table = doc.add_table(rows=rows, cols=cols)
            for r_idx, row in enumerate(data):
                for c_idx in range(cols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    table.cell(
                        r_idx, c_idx).text = "" if val is None else str(val)
            doc.save(self.file_path)
            return True
        except Exception:
            return False

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        try:
            key = str(alignment).strip().lower()
        except Exception:
            key = "left"
        mapping = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return mapping.get(key, WD_ALIGN_PARAGRAPH.LEFT)

    def _load_or_create_document(self):
        if os.path.exists(self.file_path):
            return Document(self.file_path)
        return Document()
