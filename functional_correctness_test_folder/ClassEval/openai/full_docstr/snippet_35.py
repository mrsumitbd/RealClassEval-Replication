
import os
from pathlib import Path
from typing import List

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as exc:
    raise ImportError(
        "python-docx is required for DocFileHandler. Install it via pip install python-docx") from exc


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path: str):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = Path(file_path)

    def read_text(self) -> str:
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        if not self.file_path.exists():
            return ""

        try:
            doc = Document(str(self.file_path))
            paragraphs = [para.text for para in doc.paragraphs]
            return "\n".join(paragraphs)
        except Exception:
            return ""

    def write_text(self, content: str, font_size: int = 12, alignment: str = 'left') -> bool:
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            doc = Document()
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.size = font_size * 1000  # python-docx uses 1/1000 pt
            p.alignment = self._get_alignment_value(alignment)
            doc.save(str(self.file_path))
            return True
        except Exception:
            return False

    def add_heading(self, heading: str, level: int = 1) -> bool:
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            if self.file_path.exists():
                doc = Document(str(self.file_path))
            else:
                doc = Document()
            doc.add_heading(heading, level=level)
            doc.save(str(self.file_path))
            return True
        except Exception:
            return False

    def add_table(self, data: List[List[str]]) -> bool:
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        if not data or not all(isinstance(row, list) for row in data):
            return False

        try:
            if self.file_path.exists():
                doc = Document(str(self.file_path))
            else:
                doc = Document()

            rows = len(data)
            cols = max(len(row) for row in data)
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'

            for i, row_data in enumerate(data):
                for j, cell_text in enumerate(row_data):
                    table.cell(i, j).text = str(cell_text)

            doc.save(str(self.file_path))
            return True
        except Exception:
            return False

    def _get_alignment_value(self, alignment: str) -> int:
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        mapping = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
        }
        return mapping.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)
