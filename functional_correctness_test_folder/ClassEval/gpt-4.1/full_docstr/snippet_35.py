
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        if not os.path.exists(self.file_path):
            return ""
        try:
            doc = Document(self.file_path)
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            return "\n".join(text)
        except Exception:
            return ""

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            doc = Document()
            para = doc.add_paragraph(content)
            para.alignment = self._get_alignment_value(alignment)
            for run in para.runs:
                run.font.size = Pt(font_size)
            doc.save(self.file_path)
            return True
        except Exception:
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            if os.path.exists(self.file_path):
                doc = Document(self.file_path)
            else:
                doc = Document()
            doc.add_heading(heading, level=level)
            doc.save(self.file_path)
            return True
        except Exception:
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            if not data or not isinstance(data, list) or not all(isinstance(row, list) for row in data):
                return False
            if os.path.exists(self.file_path):
                doc = Document(self.file_path)
            else:
                doc = Document()
            rows = len(data)
            cols = max(len(row) for row in data)
            table = doc.add_table(rows=rows, cols=cols)
            for i, row in enumerate(data):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = str(cell)
            doc.save(self.file_path)
            return True
        except Exception:
            return False

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        alignment = alignment.lower()
        if alignment == 'center':
            return WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            return WD_ALIGN_PARAGRAPH.RIGHT
        else:
            return WD_ALIGN_PARAGRAPH.LEFT
